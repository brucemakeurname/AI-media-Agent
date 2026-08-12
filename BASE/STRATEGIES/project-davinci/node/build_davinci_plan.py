from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = Path(__file__).with_name("02-Project-DaVinci-August-2026.docx")

NAVY = "102A43"
BLUE = "1570EF"
TEAL = "0E9384"
INK = "243B53"
MUTED = "627D98"
PALE_BLUE = "EAF2FF"
PALE_TEAL = "E6F7F5"
PALE_GRAY = "F5F7FA"
WHITE = "FFFFFF"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)

def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)

def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def add_text(paragraph, text, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return run

def set_paragraph(paragraph, before=0, after=6, line=1.15, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment

def add_bullet(doc, text, level=0, color=INK):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.55 + 0.45 * level)
    paragraph.paragraph_format.first_line_indent = Cm(-0.24)
    set_paragraph(paragraph, after=3, line=1.1)
    add_text(paragraph, text, size=10.2, color=color)
    return paragraph

def add_section_heading(doc, number, title, subtitle=None):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=12, after=4, line=1.0)
    number_run = add_text(paragraph, f"{number}. ", size=17, color=BLUE, bold=True)
    number_run.font.name = "Arial"
    add_text(paragraph, title, size=17, color=NAVY, bold=True)
    if subtitle:
        caption = doc.add_paragraph()
        set_paragraph(caption, after=9, line=1.05)
        add_text(caption, subtitle, size=9.5, color=MUTED, italic=True)

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=8, after=3, line=1.0)
    add_text(paragraph, text, size=11.5, color=TEAL, bold=True)

def style_table(table, widths=None):
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if widths:
                cell.width = Cm(widths[cell_index])
            set_cell_border(
                cell,
                top={"val": "single", "sz": "5", "color": "D9E2EC"},
                bottom={"val": "single", "sz": "5", "color": "D9E2EC"},
                left={"val": "single", "sz": "5", "color": "D9E2EC"},
                right={"val": "single", "sz": "5", "color": "D9E2EC"},
            )
            for paragraph in cell.paragraphs:
                set_paragraph(paragraph, after=2, line=1.08)
                for run in paragraph.runs:
                    set_run(run, size=8.8, color=INK)
            if row_index == 0:
                shade(cell, NAVY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run(run, size=8.8, color=WHITE, bold=True)
            elif row_index % 2 == 0:
                shade(cell, PALE_GRAY)
    if table.rows:
        set_repeat_table_header(table.rows[0])

def fill_cell(cell, lines, header=False):
    cell.text = ""
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_paragraph(paragraph, after=2, line=1.06)
        add_text(paragraph, line, size=8.8, color=WHITE if header else INK, bold=header)

def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    style = styles["List Bullet"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)

def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    header = section.header.paragraphs[0]
    set_paragraph(header, after=0, line=1.0)
    add_text(header, "ULTIMATESUP  /  AI MEDIA", size=8.5, color=TEAL, bold=True)
    add_text(header, "     PROJECT DAVINCI", size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    set_paragraph(footer, after=0, line=1.0)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(footer, "Internal working plan  |  ", size=8.5, color=MUTED)
    add_text(footer, "Page ", size=8.5, color=MUTED)
    add_page_field(footer)

def add_cover(doc):
    spacer = doc.add_paragraph()
    set_paragraph(spacer, before=20, after=14)

    eyebrow = doc.add_paragraph()
    set_paragraph(eyebrow, after=8)
    add_text(eyebrow, "AI MEDIA  •  30-DAY BUILD & PRODUCTIZATION PLAN", size=10, color=TEAL, bold=True)

    title = doc.add_paragraph()
    set_paragraph(title, after=6, line=0.92)
    add_text(title, "PROJECT DAVINCI", size=31, color=NAVY, bold=True)

    description = doc.add_paragraph()
    set_paragraph(description, after=16, line=1.18)
    add_text(
        description,
        "Xây dựng 4 workflow AI Media mới, productize để teammate tự chạy, "
        "sau đó mới tích hợp DaVinci Bot lên Lark Task.",
        size=12.2,
        color=INK,
    )

    info = doc.add_table(rows=2, cols=3)
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    entries = [
        ("NHIỆM VỤ 1", "10/08–08/09/2026"),
        ("PHẠM VI", "4 workflow mới"),
        ("NHIỆM VỤ 2", "10/09–09/11/2026"),
    ]
    for index, (label, value) in enumerate(entries):
        shade(info.cell(0, index), NAVY)
        shade(info.cell(1, index), PALE_BLUE)
        fill_cell(info.cell(0, index), [label], header=True)
        fill_cell(info.cell(1, index), [value])
    style_table(info, [5.1, 5.1, 5.1])

    note = doc.add_paragraph()
    set_paragraph(note, before=14, after=5, line=1.14)
    add_text(note, "Nguyên tắc triển khai: ", size=10.2, color=NAVY, bold=True)
    add_text(note, "workflow trước, Bot sau; mọi asset phải qua Human QA trước handoff/publish.", size=10.2, color=INK)

    marker = doc.add_paragraph()
    set_paragraph(marker, before=8, after=0)
    add_text(marker, "Phiên bản mới hoàn toàn  •  10 August 2026", size=9.2, color=MUTED, italic=True)

def add_objectives(doc):
    add_section_heading(doc, "1", "Mục tiêu", "Hệ thống workflow AI Media có thể vận hành, đo lường và bàn giao cho bất kỳ teammate nào trong team.")
    add_subheading(doc, "Nhiệm vụ 1 — Workflow Creative | 30 ngày (10/08–08/09/2026)")
    add_text(
        doc.add_paragraph(),
        "Build và productize 4 workflow AI Media mới: Video Creative (Omni), Video Clone, Summary HTML Video, Creative 2D. "
        "Mỗi workflow bao gồm setup hoàn chỉnh, input/output contract, prompt template, SOP, QA checklist, failure log và runnable sample để bất kỳ ai trong team cũng tự thực thi được.",
        size=10.4,
    )
    add_subheading(doc, "Nhiệm vụ 2 — Tích hợp Lark & DaVinci Bot | 2 tháng (10/09–09/11/2026)")
    add_text(
        doc.add_paragraph(),
        "Sau khi hoàn tất 30 ngày và chốt exit gate, chính thức đưa các workflow đã đạt chuẩn lên Lark Task. "
        "DaVinci Bot trở thành AI Media Agent của UltimateSup workspace: hỗ trợ nhận task, phân luồng, thực thi, trả status/preview/job log và chuyển asset sang Human QA. Bot không tự publish.",
        size=10.4,
    )

    outcome = doc.add_table(rows=1, cols=3)
    headers = ["KẾT QUẢ 30 NGÀY", "ĐIỀU KHÔNG LÀM", "CỔNG CHUYỂN PHA"]
    content = [
        "4 workflow productized hoàn chỉnh, có sample chạy được và có log.",
        "Không xây router/bot trên Lark khi workflow chưa đạt chuẩn.",
        "Teammate chạy theo SOP, có QA/failure data và exit-gate approval.",
    ]
    for index in range(3):
        fill_cell(outcome.cell(0, index), [headers[index], content[index]], header=False)
        shade(outcome.cell(0, index), PALE_TEAL if index != 1 else PALE_BLUE)
        cell = outcome.cell(0, index)
        for paragraph_index, paragraph in enumerate(cell.paragraphs):
            for run in paragraph.runs:
                set_run(run, size=8.7, color=TEAL if paragraph_index == 0 else INK, bold=paragraph_index == 0)
    style_table(outcome, [5.1, 5.1, 5.1])

def add_current_state(doc):
    doc.add_page_break()
    add_section_heading(doc, "2", "Thực trạng", "Hiện trạng automation media tại UltimateSup mới dừng lại ở các skill/tool rời rạc và prototype tham chiếu.")
    add_subheading(doc, "a. Những gì UltimateSup có hiện tại")
    for text in [
        "Mới chỉ có các skill/tool/prototype rời rạc; chưa tồn tại workflows creative end-to-end có thể bàn giao cho bất kỳ ai trong team.",
        "Google Flow shared tool (e7d8eab6-c1d2-4ed4-8477-effda49df52d) và các repository tham chiếu: dttstk-lab/teaclonenonelab (Read Video → Analyze → Storyboard → Obsidian), dttstk-lab/tiktok-ads-diagnostics.",
        "Kho workflow cá nhân chứa các raw workflows cần pull, kết nối tool/API và migrate vào project structure mới.",
        "Tính khả thi kỹ thuật đã được kiểm chứng. Giai đoạn 30 ngày này là xây dựng, migrate, batch-run, đo lường và cải thiện; không phải nghiên cứu khả thi.",
    ]:
        add_bullet(doc, text)

    add_subheading(doc, "b. Các workflow UltimateSup cần")
    required = doc.add_table(rows=1, cols=2)
    fill_cell(required.cell(0, 0), ["WORKFLOW MỤC TIÊU"], header=True)
    fill_cell(required.cell(0, 1), ["NĂNG LỰC CẦN BUILD MỚI SAU 7 NGÀY"], header=True)
    rows = [
        ("Video Creative (Omni)", "Biến brief/kịch bản thành video hoàn chỉnh từ script, visual direction, prompt ref, render A-roll/B-roll, voice sync đến hậu kỳ."),
        ("Video Clone", "Phân tích URL/video hiệu quả, paraphrase/re-angle script và storyboard, render biến thể mới đúng quyền nguồn và brand spec."),
        ("Summary HTML Video", "Tách bài viết/kịch bản thành frame, tạo HTML animation hyperframe, render HTML, ghép voice/audio và hậu kỳ."),
        ("Creative 2D", "Tạo hoặc clone/reverse-prompt asset 2D từ image brief và style ref, chọn lọc và retouch trước Human QA."),
    ]
    for name, scope in rows:
        cells = required.add_row().cells
        fill_cell(cells[0], [name])
        fill_cell(cells[1], [scope])
    style_table(required, [4.5, 10.8])

def add_direction(doc):
    doc.add_page_break()
    add_section_heading(doc, "3", "Định hướng", "Mỗi workflow được thiết kế với cơ chế xử lý riêng, techstack cụ thể và lớp quản trị dùng chung.")
    direction = doc.add_table(rows=1, cols=3)
    headers = ["WORKFLOW", "CƠ CHẾ HOẠT ĐỘNG", "TECHSTACK CẦN THIẾT"]
    for index, header in enumerate(headers):
        fill_cell(direction.cell(0, index), [header], header=True)
    rows = [
        (
            "Video Creative\\n(Omni)",
            "Brief/hook → visual direction → prompt/ref/shot list → render A-roll/B-roll → voice sync → hậu kỳ.",
            "Google Flow và tool render được duyệt; prompt/template; FFmpeg; voice layer có license; run log.",
        ),
        (
            "Video Clone",
            "URL/video → ingest hợp lệ → phân tích script/storyboard → paraphrase/re-angle → prompt pipeline → render → QA.",
            "Raw workflow cá nhân; prototype tham chiếu (teaclonenonelab, tiktok-ads-diagnostics); video analysis; FFmpeg; source-rights check.",
        ),
        (
            "Summary HTML\\nVideo",
            "Article/script → tách frame → HTML/CSS/JS hyperframe → render animation → voice/audio → hậu kỳ.",
            "HTML/CSS/JS; renderer được duyệt; FFmpeg; voice layer; template library.",
        ),
        (
            "Creative 2D",
            "Brief/reference → reverse prompt/style ref → render → select/retouch → QA.",
            "Image-generation tool/API được duyệt; prompt library; asset/reference store; QA checklist.",
        ),
    ]
    for row_data in rows:
        cells = direction.add_row().cells
        for index, value in enumerate(row_data):
            fill_cell(cells[index], value.split("\\n"))
    style_table(direction, [2.8, 6.2, 6.3])

    add_subheading(doc, "Lớp dùng chung & rào chắn an toàn")
    for text in [
        "Mọi workflow phải tuân thủ project structure thống nhất, input/output contract, setup guide, prompt/template, SOP, QA checklist, failure log và run log.",
        "Không thêm automation router hoặc Lark integration trong 30 ngày này; các cơ chế tự động hoá chỉ được gắn sau exit gate.",
        "AI Voice (F5-TTS, VoxCPM2, RVC, Applio) yêu cầu consent/license, intended use và quyền lưu trữ được duyệt. Private brief, customer data và thông tin sản phẩm chưa duyệt không đưa lên external tool công cộng.",
    ]:
        add_bullet(doc, text)

    callout = doc.add_table(rows=1, cols=1)
    shade(callout.cell(0, 0), PALE_BLUE)
    fill_cell(callout.cell(0, 0), [
        "TO CONFIRM",
        "Owner từng workflow package, API quota/access, nơi lưu trữ asset/job log, chính sách nguồn clone và người duyệt exit gate.",
    ])
    for paragraph_index, paragraph in enumerate(callout.cell(0, 0).paragraphs):
        for run in paragraph.runs:
            set_run(run, size=9.5, color=BLUE if paragraph_index == 0 else INK, bold=paragraph_index == 0)
    style_table(callout, [15.3])

def add_plan(doc):
    doc.add_page_break()
    add_section_heading(doc, "4", "Kế hoạch 30 ngày", "Phân bổ 30 ngày (10/08–08/09/2026) theo 4 chặng thi công, thử nghiệm batch và tinh chỉnh.")
    plan = doc.add_table(rows=1, cols=3)
    for index, header in enumerate(["THỜI GIAN", "HÀNH ĐỘNG TRIỂN KHAI CỤ THỂ", "OUTPUT / EXIT CRITERIA"]):
        fill_cell(plan.cell(0, index), [header], header=True)
    rows = [
        (
            "Ngày 1\\n10/08",
            "Chuẩn hoá project structure; kết nối các tool/API; pull raw workflows từ kho workflow cá nhân; xác nhận nơi lưu asset/log.",
            "Folder structure, access map, raw-workflow inventory, log location và draft input contract.",
        ),
        (
            "Ngày 2\\n11/08",
            "Refine hoàn chỉnh workflow Video Creative (Omni): setup, visual direction, prompt ref, render pipeline.",
            "Workflow package v0.1 + runnable sample + QA/failure log template.",
        ),
        (
            "Ngày 3\\n12/08",
            "Refine hoàn chỉnh workflow Video Clone: URL ingest, script/storyboard breakdown, paraphrase, prompt pipeline.",
            "Workflow package v0.1 + source-rights check + runnable sample.",
        ),
        (
            "Ngày 4\\n13/08",
            "Refine hoàn chỉnh workflow Summary HTML Video: article/script to frame, HTML animation, voice sync.",
            "Workflow package v0.1 + HTML render sample + QA/failure log template.",
        ),
        (
            "Ngày 5\\n14/08",
            "Refine hoàn chỉnh workflow Creative 2D: brief/ref to reverse prompt, render, select/retouch.",
            "Workflow package v0.1 + asset sample + QA/failure log template.",
        ),
        (
            "Ngày 6–7\\n15–16/08",
            "Đóng gói productize 4 workflow: chuẩn hoá setup, input/output contract, prompt/template, SOP, QA checklist, failure log và runnable sample.",
            "4 package v0.1 hoàn chỉnh sẵn sàng productize/batch-run; shared naming & log convention.",
        ),
        (
            "Ngày 8–14\\n17–23/08",
            "Mỗi ngày tạo batch nội dung cho ngày kế tiếp từ các workflows đã build; ghi nhận run log, kết quả QA và failure case.",
            "Daily batch record + preview link + status processing/done/failed + feedback data.",
        ),
        (
            "Ngày 15–21\\n24–30/08",
            "Thu thập dữ liệu từ các video/batch đã tạo; đánh giá thực trạng workflow; cải thiện logic, dataset, skill, prompt template và QA gate.",
            "Workflow package v0.2, improvement log, dataset/prompt update và failure-case update.",
        ),
        (
            "Ngày 22–30\\n31/08–08/09",
            "Teammate-run pilot: cho nhân sự ngoài dev team chạy thử theo SOP; đo success rate, cycle time, QA/rework và traceability; chốt exit gate.",
            "Báo cáo pilot Nhiệm vụ 1, 4 workflow package release candidate v1.0 và quyết định go/no-go.",
        ),
        (
            "09/09",
            "Xác nhận kết quả, owner và quyền cuối cùng trước khi khởi động Nhiệm vụ 2.",
            "Handoff record + exit-gate signoff trước khi triển khai DaVinci Bot trên Lark Task ngày 10/09.",
        ),
    ]
    for timing, focus, output in rows:
        cells = plan.add_row().cells
        fill_cell(cells[0], timing.split("\\n"))
        fill_cell(cells[1], [focus])
        fill_cell(cells[2], [output])
    style_table(plan, [2.5, 7.5, 5.3])

def add_measurement(doc):
    doc.add_page_break()
    add_section_heading(doc, "5", "Đo lường, đánh giá", "Đánh giá dựa trên dữ liệu vận hành thực tế; không dùng chỉ số lý thuyết hay tự động hoá hoàn toàn.")
    measurement = doc.add_table(rows=1, cols=4)
    for index, header in enumerate(["CHỈ SỐ ĐO LƯỜNG", "CÁCH THỨC ĐO LƯỜNG", "TẦN SUẤT", "BẰNG CHỨNG XÁC NHẬN"]):
        fill_cell(measurement.cell(0, index), [header], header=True)
    rows = [
        ("Build completeness", "4/4 workflow có đủ setup, input contract, prompt/template, SOP, QA checklist, failure log, sample.", "Ngày 7; ngày 30", "Workflow package checklist"),
        ("Run success rate", "Tỷ lệ batch run thành công tạo output đúng spec; phân loại theo prompt, render, voice, hậu kỳ, input.", "Hàng ngày", "Run log + failure log"),
        ("Processing cycle time", "Thời gian trung vị từ request hợp lệ đến preview/asset QA.", "Hàng tuần", "Timing log"),
        ("QA & rework rate", "Tỷ lệ pass Human QA ngay lần đầu, tỷ lệ phải sửa và lý do sửa.", "Hàng ngày / tuần", "QA report"),
        ("Traceability rate", "100% run có status (processing/done/failed), preview, tham số chính và job log.", "Hàng ngày", "Run log / folder log"),
        ("Pilot readiness", "Teammate ngoài nhóm dev tự thực thi được theo SOP, không cần hỗ trợ ngoài exception.", "Ngày 22–30", "Pilot report + feedback log"),
    ]
    for values in rows:
        cells = measurement.add_row().cells
        for index, value in enumerate(values):
            fill_cell(cells[index], [value])
    style_table(measurement, [3.0, 6.0, 2.6, 3.7])

    add_subheading(doc, "Điều kiện exit gate sang Nhiệm vụ 2")
    for text in [
        "4 workflow có package productized hoàn chỉnh, runnable sample và dữ liệu chạy thực tế.",
        "Teammate-run pilot cho thấy SOP đủ rõ ràng để nhân sự tự vận hành.",
        "Owner, quyền truy cập API/tool, nơi lưu job log và quy trình Human QA được duyệt chính thức.",
        "Mọi output giữ cơ chế Human QA trước handoff/publish; DaVinci Bot không tự động đăng bài.",
    ]:
        add_bullet(doc, text)

    final_note = doc.add_paragraph()
    set_paragraph(final_note, before=10, after=0, line=1.13)
    add_text(final_note, "Decision on 09/09/2026: ", size=10.2, color=NAVY, bold=True)
    add_text(final_note, "Go / No-go cho Nhiệm vụ 2 - Tích hợp DaVinci Bot lên Lark Task từ 10/09/2026.", size=10.2, color=INK)

def build():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    add_objectives(doc)
    add_current_state(doc)
    add_direction(doc)
    add_plan(doc)
    add_measurement(doc)
    doc.core_properties.title = "Project DaVinci - 30-Day Build & Productization Plan"
    doc.core_properties.subject = "AI Media workflow roadmap"
    doc.core_properties.author = "UltimateSup AI Media"
    doc.core_properties.comments = "Rebuilt from scratch on 10 August 2026"
    doc.save(OUTPUT)
    print(OUTPUT)

if __name__ == "__main__":
    build()
